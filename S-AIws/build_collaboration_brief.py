from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

OUT = r"D:\AIworkstation\S-AIws\AI工作台协作模式精简分析.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.72)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Microsoft YaHei'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.12

for name, size, color, before, after in [
    ('Heading 1', 15, '2457D6', 12, 5),
    ('Heading 2', 11.5, '1F2937', 8, 3),
]:
    st = styles[name]
    st.font.name = 'Microsoft YaHei'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p

# Running header
hp = sec.header.paragraphs[0]
hp.text = 'AI员工工作台｜协作模式分析'
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(107, 114, 128)

# Title block
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run('AI工作台协作模式精简分析')
r.font.name = 'Microsoft YaHei'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
r.font.size = Pt(23)
r.font.bold = True
r.font.color.rgb = RGBColor(31, 41, 55)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run('基于当前方案与 Claude Tag + Slack 协作机制的对照结论')
r.font.name = 'Microsoft YaHei'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(107, 114, 128)

# Lead callout
t = doc.add_table(rows=1, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
t.autofit = False
t.columns[0].width = Inches(6.74)
cell = t.cell(0,0)
cell.width = Inches(6.74)
set_cell_shading(cell, 'EEF3FF')
set_cell_margins(cell, 150, 180, 150, 180)
p = cell.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
r = p.add_run('核心定位｜')
r.bold = True
r.font.color.rgb = RGBColor(36, 87, 214)
p.add_run('个人与AI先形成工作现场；需要协同时邀请成员，会话随即升级为以人与人沟通为主、@AI协作为辅的团队工作空间。')

doc.add_heading('一、与 Claude Tag + Slack 的核心差异', level=1)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.autofit = False
widths = [Inches(1.0), Inches(2.65), Inches(3.09)]
headers = ['维度', '当前工作台', 'Claude Tag + Slack']
for i, (cell, w, text) in enumerate(zip(table.rows[0].cells, widths, headers)):
    cell.width = w
    set_cell_shading(cell, 'E8EEF8')
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9.5)
rows = [
    ('起点', '个人与AI工作，按需拉人协作', '先有团队频道，再加入AI'),
    ('默认交互', '邀请前默认AI；邀请后默认人与人', '始终以人与人沟通为主'),
    ('工作上下文', '会话、任务、文件、成果天然关联', '频道历史与外部工具数据'),
    ('闭环能力', '讨论可直接沉淀为结论、任务和成果', '通常需流转到外部系统'),
]
for vals in rows:
    cells = table.add_row().cells
    for i, (cell, w, text) in enumerate(zip(cells, widths, vals)):
        cell.width = w
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.font.size = Pt(9.2)

doc.add_heading('二、我们的核心优势', level=1)
add_bullet('协作从已有工作现场开始：新成员直接获得历史会话、文件、AI产出、任务和成果，无需重新交代背景。')
add_bullet('会话摘要不是聊天概括，而是共享工作状态：说明工作目标、最新进展、已完成、待推进和关键结论。')
add_bullet('会话天然连接执行闭环：讨论 → @AI处理 → 结论 → 任务 → 成果 → 后续流转。')

doc.add_heading('三、本期必须做', level=1)
items = [
    '模式切换可感知：首位成员加入后，明确提示会话已进入多人协作模式。',
    '共享范围可确认：邀请前告知历史消息、文件、AI产出和成果将对成员可见。',
    '共享AI上下文：所有成员使用同一个会话级AI实例，可基于他人的AI结果继续接力。',
    '交互规则统一：普通消息只发给成员，AI静默；仅有效@AI触发AI，结果全员可见并显示调用者。',
    '摘要全员共享：右侧默认打开“会话摘要”，展示工作说明、进展、完成项、待推进项和带提出人的结论；“会话目录”保持不变。',
    '邀请即更新摘要：AI生成最新摘要，发起者可修改并确认；通知展示精简摘要，新成员进入后看到完整共享摘要。',
]
for x in items:
    add_bullet(x)

doc.add_heading('四、后续重点演进', level=1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
p.add_run('讨论组织：').bold = True
p.add_run('回复链与主题线程，避免多人讨论和AI回复混杂；')
p.add_run(' 异步执行：').bold = True
p.add_run('AI后台拆解并执行长任务，完成后通知全员；')
p.add_run(' 工具联动：').bold = True
p.add_run('工具调用过程、写操作确认和结果回写全员可见；')
p.add_run(' 记忆边界：').bold = True
p.add_run('默认只读取当前会话及其已关联、已授权内容；')
p.add_run(' 管理闭环：').bold = True
p.add_run('待推进项转任务、结论确认、成员退出与移除、摘要版本历史。')

doc.add_heading('一句话结论', level=1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run('我们的方向不是“在聊天工具里加入AI”，而是把个人智能工作现场升级为团队智能协作空间：人主导沟通，AI按需介入，过程最终沉淀为可执行的任务与成果。')
r.bold = True
r.font.color.rgb = RGBColor(31, 58, 95)

# Footer
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('产品方案简报 · 2026年7月')
r.font.name = 'Microsoft YaHei'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(107, 114, 128)

doc.save(OUT)
print(OUT)
